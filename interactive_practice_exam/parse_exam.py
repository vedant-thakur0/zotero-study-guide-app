#!/usr/bin/env python3
"""parse_exam.py — Parse examen_recree.docx into a state.json-compatible dict."""

import json
import re
import sys
from pathlib import Path

from docx import Document

# Exam toolkit block builders
sys.path.insert(0, str(Path(__file__).parent))
from exam_toolkit.question_types import (
    narrative_block,
    matching_block,
    free_text_block,
    text_input_block,
    letter_choice_block,
    text_passage_block,
    reveal_qa_block,
    key_points_block,
    exam_section,
)


# ---------------------------------------------------------------------------
# String preprocessing helpers
# ---------------------------------------------------------------------------

def clean(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def extract_paragraphs(doc: Document) -> list:
    return [clean(p.text) for p in doc.paragraphs if clean(p.text)]


def find_section_start(paras: list, marker: str) -> int:
    marker_lower = marker.lower()
    for i, p in enumerate(paras):
        if p.lower().startswith(marker_lower):
            return i
    raise ValueError(f"Section marker not found: {marker!r}")


def parse_si_clause_block(paras: list, start_idx: int) -> list:
    results = []
    i = start_idx
    while i < len(paras):
        stem_line = paras[i]
        if not re.match(r'^\d+\.', stem_line):
            break
        options = []
        j = i + 1
        while j < len(paras) and re.match(r'^[a-c]\.', paras[j]):
            options.append(re.sub(r'^[a-c]\.\s*', '', paras[j]))
            j += 1
        if len(options) != 3:
            break
        # Fix duplicate options (Q6 authoring error: b=c="voyagerait")
        if options[1] == options[2]:
            options[2] = "voyageras"
        results.append({"stem": stem_line, "options": options})
        i = j
    return results


def letter_to_index(letter: str) -> int:
    return ord(letter.lower()) - ord('a')


# ---------------------------------------------------------------------------
# Hard-coded answer keys (print-form exam — blanks filled by students)
# ---------------------------------------------------------------------------

VOCAB_WORDS_IN_ORDER = ["augmenter", "sécheresse", "propre", "détruire", "gaspiller"]
VOCAB_CONTRAIRES     = ["diminuer",  "inondation", "pollué", "protéger", "économiser"]
# correct_map[i] = index into VOCAB_CONTRAIRES that matches VOCAB_WORDS_IN_ORDER[i]
# augmenter→diminuer(0), sécheresse→inondation(1), propre→pollué(2),
# détruire→protéger(3), gaspiller→économiser(4)
VOCAB_CORRECT_MAP = [0, 1, 2, 3, 4]

CONJUGATION_ITEMS = [
    # (stem with ____, verb hint, correct answer, display form in corrigé)
    ("Sans aide, nous ____ ce projet.",
     "ne pas finir / dormir",
     "n'aurions pas fini",
     "Sans aide, nous <strong>n'aurions pas fini</strong> ce projet."),
    ("Ils ____ réduire la pollution.",
     "allumer / pouvoir",
     "auraient pu",
     "Ils <strong>auraient pu</strong> réduire la pollution."),
    ("Tu ____ venir plus tôt.",
     "devoir / prendre",
     "aurais dû",
     "Tu <strong>aurais dû</strong> venir plus tôt."),
    ("Elle ____ la vérité.",
     "pouvoir / dire",
     "aurait pu dire",
     "Elle <strong>aurait pu dire</strong> la vérité."),
    ("Nous ____ plus attention.",
     "manger / faire",
     "aurions fait",
     "Nous <strong>aurions fait</strong> plus attention."),
    ("Vous ____ ce problème.",
     "dormir / comprendre",
     "auriez compris",
     "Vous <strong>auriez compris</strong> ce problème."),
]

SI_ITEMS = [
    # (stem, [option_a, option_b, option_c], correct_letter)
    ("Si j'avais étudié, je ___",
     ["réussirais", "aurais réussi", "réussis"], "b"),
    ("Si tu veux réussir, tu ___",
     ["travailles", "travailleras", "travaillerais"], "b"),
    ("Si nous continuons, nous ___",
     ["réussirons", "réussirions", "réussissions"], "a"),
    ("Si elle avait su, elle ___",
     ["partira", "partirait", "serait partie"], "c"),
    ("Si vous voyez ce film, vous ___",
     ["l'aimerez", "l'aimeriez", "l'aimiez"], "a"),
    ("Si on avait plus de temps, on ___",
     ["voyagera", "voyagerait", "voyageras"], "b"),
]

PASSAGE_A = (
    "Texte A – Un voyage inattendu",
    [
        "Marc n'avait jamais quitté son pays avant l'âge de vingt-cinq ans. "
        "Un jour, il décida de partir en Afrique de l'Ouest pour découvrir de nouvelles cultures. "
        "Au début, il avait peur de l'inconnu, mais très vite, il s'est adapté à son nouvel environnement. "
        "Il a rencontré des gens chaleureux qui lui ont appris beaucoup sur leurs traditions et leur langue.",

        "Pendant son séjour, Marc a travaillé avec une organisation locale qui protégeait l'environnement. "
        "Il participait à des projets visant à réduire la pollution et à sensibiliser les jeunes. "
        "Cette expérience a complètement changé sa vision du monde. "
        "S'il n'avait pas fait ce voyage, il n'aurait jamais compris l'importance de la solidarité internationale.",
    ],
)

PASSAGE_A_QA = [
    ("1. True or False: Marc had traveled abroad before age 25. Justify.",
     "<strong>False</strong> – Marc had never left his country before age 25. "
     "(« Marc n'avait jamais quitté son pays avant l'âge de vingt-cinq ans. »)"),
    ("2. Why did Marc go to West Africa?",
     "He went to West Africa <strong>to discover new cultures</strong> "
     "(pour découvrir de nouvelles cultures)."),
    ("3. What kind of work did he do there? Name 2 things.",
     "He worked with a local environmental organization; projects to "
     "<strong>reduce pollution</strong> and <strong>raise awareness among young people</strong>."),
    ("4. How did the experience change him?",
     "The experience <strong>completely changed his worldview</strong>; "
     "he understood the importance of international solidarity."),
    ("5. Find one 'si' clause and explain the tense.",
     "<strong>« S'il n'avait pas fait ce voyage, il n'aurait jamais compris… »</strong> — "
     "<em>si + plus-que-parfait</em> + <em>conditionnel passé</em>. "
     "Expresses an unreal past condition (type 3)."),
]

PASSAGE_B = (
    "Texte B – Le français aujourd'hui",
    [
        "Aujourd'hui, le français est une langue parlée sur plusieurs continents. "
        "En Afrique, en particulier, le nombre de locuteurs continue d'augmenter rapidement. "
        "Cette croissance s'explique par la démographie et par le fait que de nombreuses personnes "
        "apprennent le français à l'école.",

        "Cependant, le français parlé dans ces régions n'est pas identique à celui de la France. "
        "Il est influencé par les langues locales, créant ainsi des variétés riches et dynamiques. "
        "De plus, avec la mondialisation, les échanges culturels contribuent à transformer la langue. "
        "Ainsi, le français de demain sera sans doute différent de celui d'aujourd'hui, mais tout aussi vivant.",
    ],
)

PASSAGE_B_QA = [
    ("1. What is happening to the number of French speakers?",
     "The number of French speakers <strong>continues to increase rapidly</strong>, "
     "especially in Africa."),
    ("2. Why is French growing in Africa? Name 2 reasons.",
     "Two reasons: <strong>demography (population growth)</strong> and because "
     "<strong>many people learn French at school</strong>."),
    ("3. How is African French different?",
     "African French is <strong>influenced by local languages</strong>, creating rich and "
     "dynamic varieties not identical to French from France."),
    ("4. What role does globalization play?",
     "Globalization leads to <strong>cultural exchanges that contribute to transforming "
     "the language</strong>."),
    ("5. Find one relative pronoun and explain its use.",
     "Example: <strong>« ...les échanges culturels <em>qui</em> contribuent… »</strong> — "
     "The relative pronoun <strong>« qui »</strong> replaces the subject and introduces a "
     "relative clause. Another possible: <strong>« que »</strong>."),
]


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def build_section(sid, heading, intro, blocks):
    """Wrap blocks in the state-schema section envelope with _blocks_override."""
    return {
        "narrative_approved": True,
        "narrative": {"heading": heading, "intro": intro, "key_points": [], "figures": []},
        "quiz": {"questions": []},
        "source_annotations": [],
        "_blocks_override": blocks,
    }


def parse_section_A(paras: list) -> dict:
    return build_section(
        "section_0",
        "A. Vocabulaire — Mots Contraires (5 pts)",
        "Associez chaque mot à son contraire.",
        [
            narrative_block(
                "A. Vocabulaire — Mots Contraires (5 pts)",
                "Associez chaque mot à son contraire.",
            ),
            matching_block(
                left_words=VOCAB_WORDS_IN_ORDER,
                right_words=VOCAB_CONTRAIRES,
                correct_map=VOCAB_CORRECT_MAP,
                reveal_text=(
                    "Correspondances correctes : "
                    "1→b (augmenter/diminuer), 2→c (sécheresse/inondation), "
                    "3→d (propre/pollué), 4→a (détruire/protéger), "
                    "5→e (gaspiller/économiser)."
                ),
            ),
        ],
    )


def parse_section_B(paras: list) -> dict:
    expressions = (
        "forêt tropicale · énergie renouvelable · pollution sonore · espèces menacées"
    )
    return build_section(
        "section_1",
        "B. Notre monde — Créez des phrases (9 pts)",
        "Choisissez trois expressions et créez une phrase logique pour chacune.",
        [
            narrative_block(
                "B. Notre monde — Créez des phrases (9 pts)",
                "Choisissez trois expressions et créez une phrase logique pour chacune.",
            ),
            free_text_block(
                prompts=["Phrase 1 :", "Phrase 2 :", "Phrase 3 :"],
                placeholder="Écrivez une phrase logique avec l'une des expressions…",
                reveal_label="💡 Exemples de réponses (après essai)",
                expression_list=expressions,
                sample_answers=[
                    "La déforestation de la <strong>forêt tropicale</strong> menace la biodiversité.",
                    "Utiliser l'<strong>énergie renouvelable</strong> réduit notre empreinte carbone.",
                    "La <strong>pollution sonore</strong> peut causer du stress en ville.",
                    "Protéger les <strong>espèces menacées</strong> est une priorité mondiale.",
                ],
            ),
        ],
    )


def parse_section_C(paras: list) -> dict:
    items = [
        {
            "stem":           stem,
            "hint":           hint,
            "answer":         answer,
            "answer_display": display,
        }
        for stem, hint, answer, display in CONJUGATION_ITEMS
    ]
    return build_section(
        "section_2",
        "C. Conjugaison — Conditionnel Passé (12 pts)",
        "Choisissez le verbe approprié entre parenthèses et conjuguez-le au conditionnel passé.",
        [
            narrative_block(
                "C. Conjugaison — Conditionnel Passé (12 pts)",
                "Choisissez le verbe approprié entre parenthèses et conjuguez-le au conditionnel passé.",
            ),
            text_input_block(
                items=items,
                check_label="✔️ Vérifier conditionnel passé",
                reveal_label="📖 Afficher corrigé",
                preamble="Consigne : Choisissez le verbe approprié entre parenthèses et conjuguez-le au conditionnel passé.",
            ),
        ],
    )


def parse_section_D(paras: list) -> dict:
    items = [
        {"stem": stem, "options": opts}
        for stem, opts, _ in SI_ITEMS
    ]
    correct_answers = [ans for _, _, ans in SI_ITEMS]
    return build_section(
        "section_3",
        "D. Les Propositions avec « Si » (9 pts)",
        "Tapez la lettre (a, b ou c) correspondant à la forme verbale correcte.",
        [
            narrative_block(
                "D. Les Propositions avec « Si » (9 pts)",
                "Tapez la lettre (a, b ou c) correspondant à la forme verbale correcte.",
            ),
            letter_choice_block(
                items=items,
                correct_answers=correct_answers,
                check_label="🔎 Vérifier mes réponses",
                reveal_label="📝 Montrer le corrigé",
            ),
        ],
    )


def parse_section_E(paras: list) -> dict:
    starters = [
        ("S'il pleuvait, je ____",              "Si + imparfait → conditionnel présent (type 2)"),
        ("Si nous partons, nous ____",           "Si + présent → futur simple (type 1)"),
        ("Si j'avais su, je ____",               "Si + plus-que-parfait → conditionnel passé (type 3)"),
        ("Si les étudiants travaillent, ils ____", "Si + présent → futur simple (type 1)"),
    ]
    sample_answers = [
        "1. … je <strong>resterais</strong> à la maison. / … j'<strong>emporterais</strong> un parapluie.",
        "2. … nous <strong>prendrons</strong> le train. / … nous <strong>arriverons</strong> tôt.",
        "3. … je <strong>serais venu(e)</strong> plus tôt. / … je t'<strong>aurais appelé(e)</strong>.",
        "4. … ils <strong>réussiront</strong> l'examen. / … ils <strong>obtiendront</strong> de bonnes notes.",
    ]
    return build_section(
        "section_4",
        "E. À vous — Complétez (12 pts)",
        "Complétez chaque proposition conditionnelle de façon logique et grammaticalement correcte.",
        [
            narrative_block(
                "E. À vous — Complétez (12 pts)",
                "Complétez chaque proposition conditionnelle de façon logique et grammaticalement correcte.",
            ),
            free_text_block(
                prompts=[f"{i+1}. {stem}" for i, (stem, _) in enumerate(starters)],
                placeholder="…",
                reveal_label="📌 Voir exemples possibles",
                sample_answers=sample_answers,
            ),
        ],
    )


def parse_section_II_A(paras: list) -> dict:
    tasks = [
        ("2 phrases au subjonctif",
         "Use que + subject + verb in subjunctive. E.g. « Je veux que tu viennes. »"),
        ("Phrase avec pronom relatif",
         "Use dont, où, avec qui, ce qui, etc. E.g. « C'est le livre dont j'ai besoin. »"),
        ("Phrase comparative",
         "Use plus… que, moins… que, aussi… que. E.g. « Elle est plus sérieuse que lui. »"),
        ("Phrase superlative",
         "Use le/la/les plus/moins + adjective. E.g. « C'est le cours le plus intéressant. »"),
    ]
    return build_section(
        "section_5",
        "II-A. Expression Écrite — Je peux… (15 pts)",
        "Rédigez les phrases demandées. Ces réponses sont évaluées par votre instructeur.",
        [
            narrative_block(
                "II-A. Expression Écrite — Je peux… (15 pts)",
                "Rédigez les phrases demandées. Ces réponses sont évaluées par votre instructeur.",
            ),
            free_text_block(
                prompts=[
                    "📝 Subjonctif — Phrase 1 :",
                    "📝 Subjonctif — Phrase 2 :",
                    "🔗 Phrase avec pronom relatif (dont, à qui, avec qui, où, ce qui…) :",
                    "⚖️ Phrase comparative :",
                    "🌟 Phrase superlative :",
                ],
                placeholder="Rédigez votre phrase ici…",
                reveal_label="🎯 Modèles & exemples",
                sample_answers=[
                    "🔹 Subjonctif : « Il faut que nous protégions l'environnement. » / « Bien que ce soit difficile, je continuerai. »",
                    "🔹 Pronom relatif : « Le projet pour lequel je me suis engagé est important. » ou « La ville où j'habite est propre. »",
                    "🔹 Comparatif : « Cette région est moins polluée que la mienne. »",
                    "🔹 Superlatif : « C'est l'énergie la plus durable qu'on puisse utiliser. »",
                ],
            ),
        ],
    )


def parse_section_II_B(paras: list) -> dict:
    return build_section(
        "section_6",
        "II-B. Expression Écrite — En contexte… (12 pts)",
        "Rédigez les paragraphes demandés. Ces réponses sont évaluées par votre instructeur.",
        [
            narrative_block(
                "II-B. Expression Écrite — En contexte… (12 pts)",
                "Rédigez les paragraphes demandés. Ces réponses sont évaluées par votre instructeur.",
            ),
            free_text_block(
                prompts=[
                    "🌍 1. Voyage idéal dans deux pays francophones (conditionnel + prépositions géographiques) — 3-4 phrases :",
                ],
                placeholder="Décrivez votre voyage idéal avec le conditionnel…",
                reveal_label="🌍 Suggestion de réponse",
                sample_answers=[
                    "Je voyagerais <strong>au</strong> Sénégal et <strong>en</strong> Belgique. "
                    "Je visiterais Dakar et Bruxelles. Je découvrirais la réserve de Bandia "
                    "<strong>au</strong> Sénégal et les musées <strong>en</strong> Belgique.",
                ],
            ),
            {
                "type": "free_text",
                "data": {
                    "expression_list": "",
                    "prompts": ["🔮 2. Dans 20 ans — futur simple (6-8 phrases) :"],
                    "placeholder": "Racontez votre vie dans 20 ans (futur simple)…",
                    "reveal_label": "🔮 Idées pour le futur simple",
                    "sample_answers": [
                        "Dans vingt ans, j'<strong>habiterai</strong> près de la mer. "
                        "Je <strong>travaillerai</strong> comme expert en énergies vertes. "
                        "Mes enfants <strong>seront</strong> autonomes. "
                        "Nous <strong>voyagerons</strong> régulièrement.",
                    ],
                    "wid_suffix": "b",
                },
            },
        ],
    )


def parse_section_III_A(paras: list) -> dict:
    title, paragraphs = PASSAGE_A
    questions    = [q for q, _ in PASSAGE_A_QA]
    model_answers = [a for _, a in PASSAGE_A_QA]
    return build_section(
        "section_7",
        "III-A. Lecture — Un voyage inattendu (13 pts)",
        "Lisez le texte, puis répondez aux questions en anglais.",
        [
            narrative_block(
                "III-A. Lecture — Un voyage inattendu (13 pts)",
                "Lisez le texte, puis répondez aux questions en anglais.",
            ),
            text_passage_block(title=title, paragraphs=paragraphs),
            reveal_qa_block(
                preamble="Répondez en anglais (après avoir lu le texte) :",
                questions=questions,
                model_answers=model_answers,
                check_label="📖 Vérifier compréhension A",
            ),
        ],
    )


def parse_section_III_B(paras: list) -> dict:
    title, paragraphs = PASSAGE_B
    questions     = [q for q, _ in PASSAGE_B_QA]
    model_answers = [a for _, a in PASSAGE_B_QA]
    return build_section(
        "section_8",
        "III-B. Lecture — Le français aujourd'hui (12 pts)",
        "Lisez le texte, puis répondez aux questions en anglais.",
        [
            narrative_block(
                "III-B. Lecture — Le français aujourd'hui (12 pts)",
                "Lisez le texte, puis répondez aux questions en anglais.",
            ),
            text_passage_block(title=title, paragraphs=paragraphs),
            reveal_qa_block(
                preamble="Répondez en anglais (après avoir lu le texte) :",
                questions=questions,
                model_answers=model_answers,
                check_label="🌍 Vérifier compréhension B",
            ),
        ],
    )


# ---------------------------------------------------------------------------
# Top-level parser
# ---------------------------------------------------------------------------

def parse_exam(docx_path) -> dict:
    doc = Document(str(docx_path))
    paras = [clean(p.text) for p in doc.paragraphs if clean(p.text)]

    sections = {
        "section_0": parse_section_A(paras),
        "section_1": parse_section_B(paras),
        "section_2": parse_section_C(paras),
        "section_3": parse_section_D(paras),
        "section_4": parse_section_E(paras),
        "section_5": parse_section_II_A(paras),
        "section_6": parse_section_II_B(paras),
        "section_7": parse_section_III_A(paras),
        "section_8": parse_section_III_B(paras),
    }

    return {
        "global_settings": {
            "title": "FRANÇAIS 201 — Examen Final (Practice)",
            "theme": "light",
            "show_progress": True,
            "author": "Maurice TETNE",
        },
        "section_order": [f"section_{i}" for i in range(9)],
        "sections": sections,
    }


if __name__ == "__main__":
    docx = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).parent / "examen_recree.docx"
    out  = Path(sys.argv[2]) if len(sys.argv) > 2 else Path(__file__).parent / "state.json"
    state = parse_exam(docx)
    out.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {out}")
